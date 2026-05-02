"""Resume text extraction (PDF / DOCX / TXT) + LLM-based field extraction + local embeddings."""

from __future__ import annotations

import io
import json
import structlog
import re
import unicodedata
from dataclasses import dataclass
from functools import lru_cache

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.llm_provider import JSON_RESPONSE, get_llm_provider

log = structlog.get_logger()

# Validation thresholds
MIN_WORDS = 30
MAX_TEXT_CHARS = 50_000
LLM_TEXT_LIMIT = 12_000
EMBED_TEXT_LIMIT = 8_000


class ResumeParseError(ValueError):
    """Raised when a resume cannot be parsed for a known reason (encrypted, empty, etc.)."""


@dataclass
class ExtractedText:
    text: str
    word_count: int
    page_count: int  # 0 for non-paginated formats
    quality: str  # "good" | "low" | "empty"


def extract_text(filename: str, data: bytes) -> ExtractedText:
    """Extract, normalize, and validate text from a resume file.

    Raises ResumeParseError for empty/encrypted/corrupt files.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        raw, page_count = _extract_pdf(data)
    elif name.endswith(".docx"):
        raw = _extract_docx(data)
        page_count = 0
    elif name.endswith(".txt"):
        raw = data.decode("utf-8", errors="ignore")
        page_count = 0
    else:
        raise ResumeParseError(
            f"Unsupported file type: {filename}. Please upload a PDF, DOCX, or TXT."
        )

    cleaned = _clean_text(raw)
    word_count = len(cleaned.split())

    if word_count == 0:
        raise ResumeParseError(
            "We couldn't find any text in this file. "
            "If it's a scanned PDF, please upload a text-based version."
        )

    if word_count < MIN_WORDS:
        quality = "low"
    else:
        quality = "good"

    return ExtractedText(
        text=cleaned[:MAX_TEXT_CHARS],
        word_count=word_count,
        page_count=page_count,
        quality=quality,
    )


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as e:
        raise ResumeParseError(f"This PDF is corrupt or unreadable: {e}") from e

    if reader.is_encrypted:
        try:
            ok = reader.decrypt("")
            if not ok:
                raise ResumeParseError("This PDF is password-protected.")
        except Exception as e:
            raise ResumeParseError(f"This PDF is password-protected: {e}") from e

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # individual page failure shouldn't kill the whole resume
            log.warning("PDF page extraction failed: %s", e)
            parts.append("")
    return "\n\n".join(parts), len(reader.pages)


def _extract_docx(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ResumeParseError(f"This DOCX is corrupt or unreadable: {e}") from e

    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    # Many resumes use tables for layout — extract those too.
    for table in doc.tables:
        for row in table.rows:
            row_parts = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if row_parts:
                parts.append("  ".join(row_parts))
    return "\n".join(parts)


# Common PDF artifacts: bullet-point glyphs, ligatures, soft hyphens
_LIGATURES = {
    "ﬀ": "ff", "ﬁ": "fi", "ﬂ": "fl",
    "ﬃ": "ffi", "ﬄ": "ffl",
    "–": "-", "—": "-", "•": "- ",
    " ": " ", "​": "", "­": "",
}


def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalize unicode (compose accents, etc.)
    text = unicodedata.normalize("NFKC", text)
    # Replace common ligatures / typographic chars
    for src, dst in _LIGATURES.items():
        text = text.replace(src, dst)
    # Collapse runs of inline whitespace, but preserve newlines
    text = re.sub(r"[ \t]+", " ", text)
    # Trim each line
    text = "\n".join(line.strip() for line in text.splitlines())
    # Collapse 3+ newlines into 2 (preserves paragraph breaks)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------- LLM field extraction ----------

PARSE_SYSTEM = (
    "You are a precise resume parser. Given raw resume text, extract structured fields. "
    "Return STRICT JSON matching the requested schema. "
    "Do NOT invent information not present in the text. "
    "If a field is missing, omit it or use null/empty list. "
    "Preserve the candidate's original wording for highlights and summaries."
)

PARSE_USER_TEMPLATE = """Extract the following fields and return JSON only:
{{
  "full_name": string|null,
  "title": string|null,
  "summary": string|null,
  "skills": string[],
  "experience": [
    {{"company": string, "role": string, "start": string|null, "end": string|null, "highlights": string[]}}
  ],
  "education": [
    {{"institution": string, "degree": string|null, "year": string|null}}
  ],
  "projects": [{{"name": string, "description": string|null, "tech": string[]}}],
  "contact": {{"email": string|null, "phone": string|null, "location": string|null, "links": string[]}}
}}

Resume text:
\"\"\"
{text}
\"\"\""""


def _validate_parsed(parsed: dict) -> dict:
    """Ensure parsed payload conforms to the expected shape; backfill missing keys."""
    if not isinstance(parsed, dict):
        return {}
    out: dict = {
        "full_name": parsed.get("full_name") or None,
        "title": parsed.get("title") or None,
        "summary": parsed.get("summary") or None,
        "skills": [s for s in (parsed.get("skills") or []) if isinstance(s, str) and s.strip()],
        "experience": [],
        "education": [],
        "projects": [],
        "contact": None,
    }
    for e in parsed.get("experience") or []:
        if isinstance(e, dict) and (e.get("company") or e.get("role")):
            out["experience"].append(
                {
                    "company": e.get("company") or "",
                    "role": e.get("role") or "",
                    "start": e.get("start"),
                    "end": e.get("end"),
                    "highlights": [
                        h for h in (e.get("highlights") or []) if isinstance(h, str) and h.strip()
                    ],
                }
            )
    for ed in parsed.get("education") or []:
        if isinstance(ed, dict) and (ed.get("institution") or ed.get("degree")):
            out["education"].append(
                {
                    "institution": ed.get("institution") or "",
                    "degree": ed.get("degree"),
                    "year": ed.get("year"),
                }
            )
    for pr in parsed.get("projects") or []:
        if isinstance(pr, dict) and pr.get("name"):
            out["projects"].append(
                {
                    "name": pr.get("name") or "",
                    "description": pr.get("description"),
                    "tech": [t for t in (pr.get("tech") or []) if isinstance(t, str) and t.strip()],
                }
            )
    contact = parsed.get("contact")
    if isinstance(contact, dict):
        out["contact"] = {
            "email": contact.get("email"),
            "phone": contact.get("phone"),
            "location": contact.get("location"),
            "links": [
                l for l in (contact.get("links") or []) if isinstance(l, str) and l.strip()
            ],
        }
    return out


async def llm_extract_fields(text: str) -> dict:
    """Call LLM, parse JSON, validate. Never raises — returns {} on failure."""
    if not text.strip():
        return {}
    provider = get_llm_provider()
    truncated = text[:LLM_TEXT_LIMIT]
    try:
        raw = await provider.chat(
            messages=[
                {"role": "system", "content": PARSE_SYSTEM},
                {"role": "user", "content": PARSE_USER_TEMPLATE.format(text=truncated)},
            ],
            response_format=JSON_RESPONSE,
            temperature=0.0,
        )
    except Exception as e:
        log.warning("LLM call failed: %s", e)
        return {}
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        log.warning("LLM returned non-JSON; falling back to empty parse")
        return {}
    return _validate_parsed(data)


@lru_cache(maxsize=1)
def _embedder():
    from sentence_transformers import SentenceTransformer

    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")


def embed_text(text: str) -> list[float] | None:
    if not text.strip():
        return None
    try:
        model = _embedder()
        vec = model.encode(text[:EMBED_TEXT_LIMIT], normalize_embeddings=True)
        return vec.tolist()
    except Exception as e:
        log.warning("Embedding failed: %s", e)
        return None
